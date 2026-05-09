"""File-to-text extraction utilities supporting PDF, DOCX, and plain text."""

import logging
from pathlib import Path

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


def extract_text(file_path: str) -> str:
    """Read a resume or job description file and return its plain text content.

    Raises FileNotFoundError if the path does not exist, and ValueError for
    unsupported file types or files that yield no extractable text.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{suffix}'. Supported: {SUPPORTED_EXTENSIONS}"
        )

    if suffix == ".pdf":
        text = _extract_pdf(path)
    elif suffix == ".docx":
        text = _extract_docx(path)
    else:
        text = path.read_text(encoding="utf-8")

    if not text.strip():
        raise ValueError(f"No text could be extracted from: {file_path}")

    logger.debug("Extracted %d chars from %s", len(text), path.name)
    return text


def _extract_pdf(path: Path) -> str:
    """Extract text from all pages of a PDF using PyPDF2."""
    try:
        import PyPDF2
    except ImportError as exc:
        raise ImportError("PyPDF2 is required for PDF parsing: pip install PyPDF2") from exc

    pages: list[str] = []
    with path.open("rb") as fh:
        reader = PyPDF2.PdfReader(fh)
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text() or ""
            if page_text.strip():
                pages.append(page_text)
            else:
                logger.warning("Page %d of %s yielded no text (may be image-based)", i + 1, path.name)

    return "\n\n".join(pages)


def _extract_docx(path: Path) -> str:
    """Extract text from paragraphs and tables in a DOCX file."""
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError("python-docx is required for DOCX parsing: pip install python-docx") from exc

    doc = Document(str(path))
    parts: list[str] = []

    for para in doc.paragraphs:
        stripped = para.text.strip()
        if stripped:
            parts.append(stripped)

    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                parts.append("  |  ".join(row_cells))

    return "\n".join(parts)
